import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Set path for diagrams
output_dir = os.path.join(os.path.dirname(__file__), "images")

# Fonts
FONT_NAME = "Arial"

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Enable RTL for the section
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

    return doc

# RTL Helpers
def make_paragraph_rtl(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._element.get_or_add_pPr()
    pBidi = OxmlElement('w:bidi')
    pBidi.set(qn('w:val'), '1')
    pPr.append(pBidi)

def add_arabic_run(p, text, bold=False, italic=False, size=11, color=None):
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
        
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rPr.append(rFonts)
    
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    return run

def add_heading_rtl(doc, text, level, space_before=12, space_after=6):
    p = doc.add_paragraph()
    make_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        size = 18
        color = RGBColor(0, 76, 153) # Dark Blue
        bold = True
    elif level == 2:
        size = 14
        color = RGBColor(0, 102, 204) # Medium Blue
        bold = True
    elif level == 3:
        size = 12
        color = RGBColor(51, 51, 51) # Dark Gray
        bold = True
    else:
        size = 11
        color = RGBColor(51, 51, 51)
        bold = True
        
    add_arabic_run(p, text, bold=bold, size=size, color=color)
    return p

def add_body_rtl(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    make_paragraph_rtl(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    add_arabic_run(p, text, bold=bold, italic=italic, size=size)
    return p

# Table Formatting
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def add_table_rtl(doc, headers, rows):
    cols_count = len(headers)
    table = doc.add_table(rows=1, cols=cols_count)
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_table_borders(table)
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        make_paragraph_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_arabic_run(p, header_text, bold=True, size=10, color=RGBColor(255, 255, 255))
        set_cell_shading(hdr_cells[i], "004C99")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
    # Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            make_paragraph_rtl(p)
            
            # Simple check for numbers
            is_num = all(c.isdigit() or c in '$.%:-/ ' for c in str(cell_value))
            if is_num:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
            add_arabic_run(p, str(cell_value), size=9.5)
            set_cell_shading(row_cells[i], "F2F8FF" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], top=100, bottom=100, left=120, right=120)
            
    # Spacing paragraph after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_image_rtl(doc, img_name, caption=""):
    img_path = os.path.join(output_dir, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(img_path, width=Inches(5.8))
        
        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_after = Pt(12)
            add_arabic_run(cap_p, caption, italic=True, size=9.5, color=RGBColor(102, 102, 102))
    else:
        print(f"Warning: Image {img_name} not found!")

def add_page_number(run):
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p_hdr = header.paragraphs[0]
        make_paragraph_rtl(p_hdr)
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_arabic_run(p_hdr, "منصة سَنَد للتكافل الطبي وإعارة الأجهزة الطبية", size=8.5, color=RGBColor(150, 150, 150))
        
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        make_paragraph_rtl(p_ftr)
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = add_arabic_run(p_ftr, "صفحة ", size=9, color=RGBColor(120, 120, 120))
        add_page_number(run)

def build_cover_page(doc):
    # Align Center for Cover Page
    for i in range(3):
        doc.add_paragraph()
        
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p1, "الجمهورية اليمنية\nوزارة التعليم الفني والتدريب المهني\nكلية المجتمع - صنعاء\nقسم تكنولوجيا المعلومات (IT)", bold=True, size=14)
    
    for i in range(4):
        doc.add_paragraph()
        
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p2, "مشروع تخرج لنيل درجة الدبلوم في تكنولوجيا المعلومات", size=13, color=RGBColor(80, 80, 80))
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p3, "\nمنصة سَنَد\nلتكافل وإعارة الأجهزة الطبية المستعملة", bold=True, size=24, color=RGBColor(0, 76, 153))
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p4, "بوابة إلكترونية لتنظيم وإدارة تداول وإعارة وتبرع الأجهزة والمستلزمات الطبية", size=11, color=RGBColor(100, 100, 100))
    
    for i in range(5):
        doc.add_paragraph()
        
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p5, "إعداد الطالب:\nفارس أنور\n\nإشراف الدكتور:\nأحمد الناشري", bold=True, size=12)
    
    for i in range(3):
        doc.add_paragraph()
        
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p6, "2025م / 1446هـ", bold=True, size=11)
    
    doc.add_page_break()

def build_preliminary_pages(doc):
    # Dedication
    add_heading_rtl(doc, "الإهداء", level=1)
    p = doc.add_paragraph()
    make_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(36)
    add_arabic_run(p, "إلى من ساروا على دروب العلم فأناروا لنا الطريق... آباؤنا وأمهاتنا الأوفياء، من وهبونا الحياة وبذلوا كل ثمين لرؤيتنا ننجح.\n\nإلى أساتذتنا الكرام في قسم تكنولوجيا المعلومات بكلية المجتمع، ممن أعطوا بلا حدود وزرعوا فينا الشغف والاجتهاد.\n\nإلى زملائنا الأعزاء، شركاء الدرب والسهر والتحصيل العلمي.\n\nوإلى وطننا الحبيب اليمن، الذي نطمح أن نسهم في بنائه وتطوير قطاعاته التكافلية والطبية من خلال ما تعلمناه.\n\nنهدي هذا العمل المتواضع.", size=12, italic=True)
    doc.add_page_break()
    
    # Acknowledgment
    add_heading_rtl(doc, "شكر وتقدير", level=1)
    add_body_rtl(doc, "الحمد لله رب العالمين، والصلاة والسلام على رسوله الكريم، معلم البشرية وهاديها، وعلى آله وصحبه الأخيار.")
    add_body_rtl(doc, "في البداية، نتوجه بالشكر الجزيل والحمد لرب العالمين الذي أنعم علينا بالصحة والعزيمة لإتمام هذا العمل التكافلي الإنساني والتقني.")
    add_body_rtl(doc, "كما يسعدنا أن نتقدم ببالغ الشكر وعظيم الامتنان لمشرف هذا المشروع الدكتور الفاضل/ أحمد الناشري، الذي لم يبخل علينا بوقته وتوجيهاته القيمة التي كان لها أكبر الأثر في توجيه مسار هذا المشروع بالشكل الأكاديمي والبرمجي الصحيح.")
    add_body_rtl(doc, "ونتوجه بالشكر والتقدير إلى عمادة كلية المجتمع بصنعاء، ورئيس قسم تكنولوجيا المعلومات وجميع أعضاء الهيئة التدريسية على كل ما قدموه لنا من دعم ومعرفة طيلة سنوات الدراسة والتحصيل العلمي.")
    add_body_rtl(doc, "وفي الختام، نشكر كل من قدم لنا نصيحة أو تشجيعاً أو ساعدنا في جمع متطلبات هذا النظام لإنجاحه وإخراجه إلى النور لخدمة المجتمع اليمني.")
    doc.add_page_break()
    
    # Table of Contents (Manual list style table)
    add_heading_rtl(doc, "الفهرس", level=1)
    toc_headers = ["الموضوع", "الصفحة"]
    toc_rows = [
        ["الإهداء", "3"],
        ["شكر وتقدير", "4"],
        ["الفهرس", "5"],
        ["ملخص المشروع", "6"],
        ["المختصرات والمصطلحات", "7"],
        ["الفصل الأول: الإجراءات المنهجية للمشروع", "8"],
        ["1-1 المقدمة", "9"],
        ["1-2 تعريف المشروع", "10"],
        ["1-3 تحديد المشكلة", "10"],
        ["1-4 تعريف المشكلة", "10"],
        ["1-5 الأهداف (الهدف العام والأهداف الخاصة)", "11"],
        ["1-6 أهمية المشروع", "11"],
        ["1-7 مدى المشروع", "12"],
        ["1-8 حدود ونطاق المشروع", "12"],
        ["1-9 منهجية تحليل وتصميم النظام", "14"],
        ["1-10 المنهجية المستخدمة (SDLC)", "14"],
        ["1-11 الخطة الزمنية لإنجاز المشروع", "16"],
        ["1-12 دراسة الجدوى (التقنية والاقتصادية والتشغيلية)", "17"],
        ["الفصل الثاني: الدراسات السابقة (خلفية نظرية)", "22"],
        ["2-1 مقدمة", "23"],
        ["2-2 منصة MedShare العالمية", "23"],
        ["2-3 منصة Freecycle لتبادل السلع", "24"],
        ["2-4 منصة ShareMed للأجهزة الطبية", "25"],
        ["2-5 مقارنة بين منصة سند والمنصات السابقة", "26"],
        ["الفصل الثالث: التحليل", "28"],
        ["3-1 المقدمة", "29"],
        ["3-2 طرق تحديد المتطلبات", "29"],
        ["3-3 هيكلية النظام", "30"],
        ["3-4 فوائد التحليل", "31"],
        ["3-5 نموذج الأحداث Event Modeling", "32"],
        ["3-6 مخططات تدفق البيانات DFD", "34"],
        ["3-7 مخطط العلاقات الكينونية ERD", "37"],
        ["3-8 لغة النمذجة الموحدة UML", "39"],
        ["الفصل الرابع: التصميم", "47"],
        ["4-1 المقدمة", "48"],
        ["4-2 المخطط الانسيابي Flowcharts", "48"],
        ["4-3 جداول قاعدة البيانات", "52"],
        ["4-4 واجهات المنصة", "62"],
        ["الفصل الخامس: الخاتمة والتوصيات", "78"],
        ["5-1 الخاتمة", "79"],
        ["5-2 التوصيات", "80"],
        ["5-3 العمل المستقبلي للتطوير", "80"],
        ["5-4 المعوقات", "81"],
        ["5-5 المراجع", "82"]
    ]
    add_table_rtl(doc, toc_headers, toc_rows)
    doc.add_page_break()
    
    # Abstract
    add_heading_rtl(doc, "ملخص المشروع", level=1)
    add_body_rtl(doc, "المشروع عبارة عن منصة رقمية تكافلية مبتكرة باسم (سَنَد)، تهدف إلى سد الفجوة بين الأسر التي تمتلك أجهزة ومستلزمات طبية فائضة مخزنة في بيوتها، وبين الفئات الفقيرة والمحتاجة من المرضى غير القادرين على توفير أو شراء هذه الأجهزة الباهظة بسبب الأوضاع المعيشية الحالية في المجتمع اليمني.")
    add_body_rtl(doc, "تعتمد آلية عمل المنصة على تقديم نظام متكامل يتكامل مع قنوات التواصل والتحقق البشري؛ حيث يقوم المتبرع (المعير) بإدراج بيانات وموقع الجهاز بدقة على خريطة تفاعلية، بينما يقوم المستفيد بتقديم طلب استعارة مرفقاً بتقرير طبي يثبت حاجته الفعلية. يقوم مشرف النظام (Admin) بمراجعة الطلب والمستندات الطبية بدقة لضمان وصول الدعم للمستحقين فعلياً ومنع أي احتكار أو تلاعب.")
    add_body_rtl(doc, "بمجرد موافقة المشرف، يتم إتاحة خيارات التواصل المباشر للمريض مع المتبرع عبر تطبيق الواتساب أو الاتصال الهاتفي، مع الحفاظ على الخصوصية وحماية البيانات الطبية التي تخزن بشكل آمن في الخادم. تم تنفيذ المنصة باستخدام تقنيات برمجية خام (Pure PHP, Vanilla JS, Custom CSS) بدون الاستعانة بأطر عمل جاهزة، لبيان وفهم أساسيات الهندسة البرمجية والتحكم التام في جودة التصميم والسرعة والحماية.")
    doc.add_page_break()
    
    # Abbreviations
    add_heading_rtl(doc, "المختصرات والمصطلحات", level=1)
    abb_headers = ["المصطلح", "التعريف باللغة الإنجليزية", "التعريف باللغة العربية"]
    abb_rows = [
        ["DFD", "Data Flow Diagram", "مخطط تدفق البيانات"],
        ["ERD", "Entity Relationship Diagram", "مخطط العلاقات الكينونية"],
        ["SDLC", "System Development Life Cycle", "دورة حياة تطوير النظم"],
        ["PHP", "Hypertext Preprocessor", "لغة برمجة خادم الويب (معالج النصوص الفائقة)"],
        ["SQL", "Structured Query Language", "لغة الاستعلام البنيوية لقواعد البيانات"],
        ["UML", "Unified Modeling Language", "لغة النمذجة الموحدة للمخططات البرمجية"],
        ["CSS", "Cascading Style Sheets", "أوراق الأنماط المتتالية لتنسيق المواقع"],
        ["HTML", "Hypertext Markup Language", "لغة ترميز النصوص الفائقة لهيكلة الويب"],
        ["JS", "JavaScript", "لغة برمجة تفاعل صفحات الويب (جافا سكريبت)"],
        ["API", "Application Programming Interface", "واجهة برمجة التطبيقات للتكامل الخارجي"],
        ["RTL", "Right to Left", "اتجاه الكتابة والتنسيق من اليمين إلى اليسار"]
    ]
    add_table_rtl(doc, abb_headers, abb_rows)
    doc.add_page_break()

def build_chapter_1(doc):
    add_heading_rtl(doc, "الفصل الأول: الإجراءات المنهجية للمشروع", level=1)
    
    add_heading_rtl(doc, "1-1 المقدمة", level=2)
    add_body_rtl(doc, "في ظل الأزمات المعاصرة والتحديات الاقتصادية المتسارعة، تواجه المجتمعات النامية وعلى رأسها المجتمع اليمني صعوبات بالغة في تلبية المتطلبات الأساسية للرعاية الصحية. فقد شهد قطاع الصحة والخدمات الطبية تحديات كبيرة انعكست سلباً على قدرة المواطن العادي في شراء الأدوية والمستلزمات الطبية التخصصية، ناهيك عن الأجهزة الطبية الضرورية مثل مكثفات الأكسجين، الكراسي المتحركة، والأسرة الطبية السريرية.")
    add_body_rtl(doc, "ومع التطور التكنولوجي الهائل، أصبح من الضروري استغلال الحلول البرمجية وتطبيقات الويب لخدمة العمل التكافلي التبرعي وتنظيم قنوات إعارة وتداول الأجهزة المستعملة. يهدف هذا البحث والعمل البرمجي إلى تقديم منصة ويب وطنية باسم (سَنَد) تعمل كحلقة وصل تقنية آمنة ومنظمة بين المتبرعين ممن لديهم أجهزة طبية خاملة لم يعودوا بحاجة إليها، والمرضى الفعليين المحتاجين لها جغرافياً وإنسانياً.")
    
    add_heading_rtl(doc, "1-2 تعريف المشروع", level=2)
    add_body_rtl(doc, "منصة (سَنَد) هي تطبيق ويب متجاوب تكافلي مجتمعي، مصمم خصيصاً ليناسب طبيعة البيئة والإنترنت في اليمن. تتيح المنصة للمستخدمين إدراج الأجهزة الطبية المستعملة المتوفرة لديهم (سواء للتبرع النهائي أو الإعارة المؤقتة) وتحديد موقعه الجغرافي الدقيق باستخدام خطوط الطول والعرض. كما يتيح للمحتاجين استعراض وتصفية هذه الأجهزة والتقدم بطلب الحصول عليها مع توثيق الحالة الطبية برفع التقارير الرسمية، لتقوم إدارة المنصة بالتدقيق والتحقق قبل تفعيل التواصل المباشر بين الطرفين.")
    
    add_heading_rtl(doc, "1-3 تحديد المشكلة", level=2)
    add_body_rtl(doc, "تكمن المشكلة في وجود آلاف العائلات التي تعاني لتوفير إيجار أو قيمة شراء جهاز تنفس أو كرسي متحرك لمرضاهم، مما يهدد حياتهم بالخطر، بينما في المقابل تحتفظ عائلات أخرى بأجهزة مماثلة خاملة في بيوتهم بعد شفاء مرضاهم دون استغلال. ولا يوجد حالياً نظام رقمي منظم محلي يربط بين الطرفين، مما يؤدي إلى هدر موارد صحية ثمينة وضياع فرص لإنقاذ الأرواح بسبب صعوبة التواصل الموثوق والجغرافي.")
    
    add_heading_rtl(doc, "1-4 تعريف المشكلة", level=2)
    add_body_rtl(doc, "يمكن بلورة المشكلة الأساسية في النقاط التالية:")
    add_body_rtl(doc, "1. عدم توافر قناة اتصال رقمية موثوقة ومخصصة لإدارة إعارة الأجهزة الطبية المستعملة محلياً.")
    add_body_rtl(doc, "2. الارتفاع الهائل في أسعار المستلزمات والأجهزة الطبية الجديدة بفعل الظروف الاقتصادية في اليمن.")
    add_body_rtl(doc, "3. هدر الموارد الطبية الصالحة للاستخدام نتيجة تخزينها بشكل خامل في المنازل بعد انتهاء الحاجة الشخصية لها.")
    add_body_rtl(doc, "4. صعوبة التحقق من حاجة المتقدمين بطلبات الأجهزة عبر منصات التواصل العام (فيسبوك، واتساب) مما يسهل الاحتكار أو إعادة البيع.")
    
    add_heading_rtl(doc, "1-5 الأهداف", level=2)
    add_heading_rtl(doc, "الهدف العام", level=3)
    add_body_rtl(doc, "تطوير وبناء منصة ويب تكافلية متكاملة باسم (سَنَد) لتنظيم وتبادل وإعارة وتبرع الأجهزة الطبية المستعملة في اليمن بشكل آمن، جغرافي، وموثوق إنسانياً وتحت إشراف رقابي.")
    
    add_heading_rtl(doc, "الأهداف الخاصة", level=3)
    add_body_rtl(doc, "1. تصميم واجهات مستخدم متجاوبة (RTL) تناسب شاشات الهواتف والحواسب وسرعات الإنترنت المحلية.")
    add_body_rtl(doc, "2. بناء قاعدة بيانات علائقية متينة وآمنة لحفظ بيانات الحسابات، الأجهزة، والصور والطلبات.")
    add_body_rtl(doc, "3. دمج خرائط جوجل (Google Maps) لتمكين المتبرعين والمستفيدين من التحديد الجغرافي الدقيق للأجهزة الطبية.")
    add_body_rtl(doc, "4. بناء نظام تحقق طبي رقابي يسمح للمشرفين بفحص التقارير والروشتات الطبية المرفوعة قبل السماح بالتواصل.")
    add_body_rtl(doc, "5. تفعيل قنوات الاتصال المباشر التفاعلي (واتساب بضغط زر واحدة / اتصال مباشر) لتسريع تسليم الأجهزة بعد القبول الإداري.")
    
    add_heading_rtl(doc, "1-6 أهمية المشروع", level=2)
    add_body_rtl(doc, "تتجلى أهمية منصة سند في جانبين أساسيين:")
    add_body_rtl(doc, "الأول: إنساني وتكافلي؛ عبر تقليل الأعباء المالية الضخمة عن عائلات المرضى الفقراء وتفعيل مبدأ الرعاية التشاركية والاقتصاد الدائري (Zero Waste) في المجال الطبي.")
    add_body_rtl(doc, "الثاني: تقني وبرمجي؛ حيث يمثل التطبيق نموذجاً تطبيقياً للويب التفاعلي الآمن باستخدام تقنيات برمجية خام (Pure PHP / MySQL / CSS / JS) مما يبرهن على قدرة المهندسين على فهم كواليس البناء وتأمين النماذج ومعالجة الملفات وحماية الخصوصية دون اعتماد أطر عمل تخفي التفاصيل.")
    
    add_heading_rtl(doc, "1-7 مدى المشروع", level=2)
    add_body_rtl(doc, "يستهدف المشروع عموم المحافظات والمديريات اليمنية، ويركز على فئات الأجهزة الطبية الأساسية (مثل أجهزة التنفس، أجهزة الحركة، الأثاث الطبي، أجهزة التشخيص والقياس). ويتضمن فترات استعارة مؤقتة محددة أو تبرع دائم بالكامل.")
    
    add_heading_rtl(doc, "1-8 حدود ونطاق المشروع", level=2)
    add_body_rtl(doc, "يتم تحديد صلاحيات وأدوار النظام عبر الأدوار الأربعة الموضحة في الجدول التالي:")
    
    scope_headers = ["الدور", "المسمى الوظيفي", "أبرز الصلاحيات البرمجية"]
    scope_rows = [
        ["الزائر", "Guest / Visitor", "تصفح سوق الأجهزة، البحث والفلترة الجغرافية حسب المحافظة والمديرية، إنشاء حساب جديد."],
        ["المتبرع / المُعير", "Donor / Lender", "إدراج جهاز جديد، رفع الصور، تحديد الموقع على الخريطة، تحديد شروط ومدة الإعارة، تصفح المستفيدين المقبولين."],
        ["المستفيد / المريض", "Beneficiary / Patient", "البحث عن جهاز، إرسال طلب استعارة، رفع مستندات التقرير الطبي، الحصول على أزرار واتساب للتواصل بعد القبول."],
        ["مدير النظام / المشرف", "Admin / Moderator", "لوحة تحكم كاملة، تفعيل أو رفض الأجهزة المضافة، فحص وتدقيق التقارير الطبية والموافقة على طلبات الاحتياج."]
    ]
    add_table_rtl(doc, scope_headers, scope_rows)
    
    add_heading_rtl(doc, "1-9 منهجية تحليل وتصميم النظام", level=2)
    add_body_rtl(doc, "اعتمد المشروع على أسلوب علمي منظم لدراسة وتخطيط وتطوير النظام لضمان الاستدامة وتقليل الأخطاء البرمجية أثناء التسليم والمناقشة.")
    
    add_heading_rtl(doc, "1-10 المنهجية المستخدمة (SDLC)", level=2)
    add_body_rtl(doc, "تم اتباع منهجية دورة حياة تطوير البرمجيات (System Development Life Cycle - SDLC) وتحديداً النموذج الشلالي (Waterfall Model) نظراً لثبات المتطلبات ووضوح أهداف التطبيق كالتالي:")
    add_body_rtl(doc, "1. التخطيط وجمع المتطلبات (Planning): تحديد المشكلة، الأهداف، والجدوى.")
    add_body_rtl(doc, "2. التحليل (Analysis): تحديد المتطلبات الوظيفية وغير الوظيفية، وبناء مخططات DFD وERD وحالات الاستخدام.")
    add_body_rtl(doc, "3. التصميم (Design): تصميم هيكل قاعدة البيانات، الواجهات الرسومية، والأنماط الانسيابية.")
    add_body_rtl(doc, "4. التطوير والترميز (Implementation): كتابة الكود البرمجي بلغة PHP وHTML/CSS/JS.")
    add_body_rtl(doc, "5. الاختبار والتحقق (Testing): فحص ثغرات المدخلات وتوافق الشاشات وأداء الاستعلامات.")
    add_body_rtl(doc, "6. التوثيق والنشر (Documentation): كتابة ملفات التوجيه والكتيب الفني للمناقشة.")
    
    add_heading_rtl(doc, "1-11 الخطة الزمنية لإنجاز المشروع", level=2)
    time_headers = ["المرحلة البرمجية", "تاريخ البدء", "تاريخ الانتهاء", "المدة المقدرة (يوم)"]
    time_rows = [
        ["جمع البيانات والتخطيط", "2025/01/01", "2025/01/20", "20"],
        ["التحليل وهيكلة المخططات DFD/ERD", "2025/01/21", "2025/02/10", "20"],
        ["تصميم قاعدة البيانات MySQL", "2025/02/11", "2025/02/25", "15"],
        ["تصميم واجهات الويب وأوراق CSS", "2025/02/26", "2025/03/12", "15"],
        ["البرمجة والتطوير (PHP Backend)", "2025/03/13", "2025/04/12", "30"],
        ["الاختبار والتصحيح وفحص الثغرات", "2025/04/13", "2025/04/28", "15"],
        ["كتابة كتيب التوثيق الفني والنشر", "2025/04/29", "2025/05/18", "20"]
    ]
    add_table_rtl(doc, time_headers, time_rows)
    
    add_heading_rtl(doc, "1-12 دراسة الجدوى", level=2)
    add_heading_rtl(doc, "الجدوى التقنية (Technical Feasibility)", level=3)
    add_body_rtl(doc, "المتطلبات المادية والبرمجية لتشغيل وتطوير منصة سند متوفرة بالكامل ولا تحتاج إلى أجهزة متخصصة باهظة كالتالي:")
    
    tech_hw_headers = ["المتطلبات المادية", "المواصفات الأدنى المطلوبة"]
    tech_hw_rows = [
        ["جهاز حاسوب للمطور / السيرفر", "Intel Core i3, 4GB RAM, 100GB Disk Space"],
        ["هاتف ذكي للمستفيد / المتبرع", "Android or iOS operating system with internet capability"],
        ["الاتصال بالإنترنت", "Modem / Mobile Data (2G/3G/4G) sufficient for browsing"]
    ]
    add_table_rtl(doc, tech_hw_headers, tech_hw_rows)
    
    tech_sw_headers = ["المتطلبات البرمجية", "الغرض منها في التطوير"]
    tech_sw_rows = [
        ["بيئة الخادم المحلي", "XAMPP / WampServer (Apache + PHP 8.2 + MySQL)"],
        ["محرر الأكواد البرمجية", "Visual Studio Code (VS Code)"],
        ["محرك قواعد البيانات", "MySQL Client (phpMyAdmin)"],
        ["لغات وتنسيقات الويب", "Pure HTML5, CSS3, ES6+ Vanilla JavaScript"],
        ["محرك لغة البرمجة", "Pure PHP (Procedural/Object-Oriented)"]
    ]
    add_table_rtl(doc, tech_sw_headers, tech_sw_rows)
    
    add_heading_rtl(doc, "الجدوى الاقتصادية (Economic Feasibility)", level=3)
    add_body_rtl(doc, "يتميز التطبيق بجدوى اقتصادية ممتازة نظراً لاعتماده الكامل على برمجيات مفتوحة المصدر (Open Source) وتطوير فردي دون تكلفة تراخيص باهظة:")
    
    eco_headers = ["بند التكلفة", "النظام التقليدي (تبرعات عشوائية)", "نظام منصة سَنَد التكافلية"]
    eco_rows = [
        ["تكلفة التراخيص البرمجية", "$500 (Commercial Platforms)", "$0 (Open Source PHP/MySQL)"],
        ["استضافة الخادم والدومين", "$0 (Manual local calls)", "$15 - $30 / Year (Basic Shared Hosting)"],
        ["تكاليف النقل والتحقق", "باهظة (لجان ميدانية وزيارات)", "$0 (تحقق رقمي عبر رفع الوثائق والتقارير)"],
        ["تكاليف التسويق والإعلانات", "مرتفعة (مطبوعات وإعلانات ممولة)", "$0 (عبر قنوات التواصل ومجموعات واتساب)"]
    ]
    add_table_rtl(doc, eco_headers, eco_rows)
    
    add_heading_rtl(doc, "الجدوى التشغيلية (Operational Feasibility)", level=3)
    add_body_rtl(doc, "المنصة قابلة للتشغيل المباشر وبسيطة جداً؛ حيث تعتمد على واجهات مستخدم مألوفة بالكامل باللغة العربية، وتختصر آلية التواصل في روابط واتساب تلقائية، مما لا يتطلب أي تدريب تقني لكبار السن أو المستخدمين العاديين.")
    
    doc.add_page_break()

def build_chapter_2(doc):
    add_heading_rtl(doc, "الفصل الثاني: الدراسات السابقة (خلفية نظرية)", level=1)
    
    add_heading_rtl(doc, "2-1 مقدمة", level=2)
    add_body_rtl(doc, "تعد دراسة الأنظمة المشابهة إقليمياً ودولياً خطوة هامة في هندسة البرمجيات، حيث تمنح المطورين فهماً واسعاً لأفضل الممارسات (Best Practices) في معالجة تدفق البيانات، وتصميم تجربة المستخدم (UX)، وتجنب العيوب البرمجية والتصميمية.")
    
    add_heading_rtl(doc, "2-2 منصة MedShare العالمية", level=2)
    add_body_rtl(doc, "تعتبر MedShare منظمة غير ربحية عالمية تجمع الأجهزة الطبية الفائضة من المستشفيات والشركات المصنعة في الدول المتقدمة وتقوم بفرزها وشحنها للمستشفيات والعيادات الطبية في الدول النامية. تتميز المنصة بالقدرة على شحن حاويات ضخمة من الأجهزة الحيوية وتدريب الطواقم الفنية على استخدامها.")
    add_body_rtl(doc, "أبرز العيوب مقارنة بسند: تركز MedShare على التبرعات المؤسسية الضخمة (بين المستشفيات والدول) ولا تتيح للأفراد والبيوت إعارة أجهزة صغيرة فيما بينهم بشكل محلي مباشر وسريع.")
    
    add_heading_rtl(doc, "2-3 منصة Freecycle لتبادل السلع", level=2)
    add_body_rtl(doc, "منصة Freecycle هي شبكة عالمية تضم ملايين الأعضاء تهدف لتشجيع إعادة تدوير المواد القديمة وتقليل النفايات عبر تبرع الأفراد بكل ما لديهم من ملابس وأثاث وأدوات مستعملة لجيرانهم مجاناً.")
    add_body_rtl(doc, "أبرز العيوب مقارنة بسند: منصة عامة جداً تفتقر تماماً للتحقق الطبي، ولا تلزم برفع تقارير طبية، مما يجعلها غير صالحة للأجهزة الطبية الحساسة التي قد تستغل لأغراض تجارية أو تباع في السوق السوداء.")
    
    add_heading_rtl(doc, "2-4 منصة ShareMed للأجهزة الطبية", level=2)
    add_body_rtl(doc, "منصة ShareMed هي تطبيق إلكتروني مخصص لتبادل وإعارة الأجهزة الطبية الفردية في بعض البلدان المتقدمة. يعتمد النظام على شركات شحن محلية تتولى نقل الأجهزة من بيت المتبرع إلى المريض بعد حجزها عبر الموقع.")
    add_body_rtl(doc, "أبرز العيوب مقارنة بسند: تعتمد كلياً على تكامل شركات الشحن الرقمية وتتطلب دفع رسوم توصيل إلكترونية عبر فيزا وماستركارد، وهو ما لا يتناسب مطلقاً مع شح بوابات الدفع الإلكتروني والظروف الحالية في اليمن.")
    
    add_heading_rtl(doc, "2-5 مقارنة بين منصة سند والمنصات السابقة", level=2)
    compare_headers = ["وجه المقارنة", "منصة MedShare", "منصة Freecycle", "منصة ShareMed", "منصة سَنَد (المقترحة)"]
    compare_rows = [
        ["الفئة المستهدفة", "مستشفيات ومؤسسات", "عامة (كل السلع)", "أفراد ومرضى", "مرضى وأفراد يمنيون"],
        ["نظام التحقق الطبي", "شديد وصارم (مؤسسي)", "لا يوجد نهائياً", "مستندات مبسطة", "إلزامي (تقارير معتمدة)"],
        ["تحديد الموقع الجغرافي", "دولي وشحن بحري", "نصوص عامة", "خرائط وتوصيل آلي", "خريطة تفاعلية محلية"],
        ["آلية التواصل", "رسمية وبيروقراطية", "بريد إلكتروني داخلي", "محادثات داخلية", "واتساب بلمسة واحدة"],
        ["ملاءمة البيئة اليمنية", "ضعيفة (تتطلب تراخيص دولية)", "غير صالحة طبياً", "صعبة (بسبب وسائل الدفع)", "ممتازة جداً (بسيطة ومباشرة)"]
    ]
    add_table_rtl(doc, compare_headers, compare_rows)
    
    doc.add_page_break()

def build_chapter_3(doc):
    add_heading_rtl(doc, "الفصل الثالث: التحليل", level=1)
    
    add_heading_rtl(doc, "3-1 المقدمة", level=2)
    add_body_rtl(doc, "تعتبر مرحلة التحليل حجر الأساس في هندسة البرمجيات؛ حيث يتركز العمل على فهم هيكلية النظام والوظائف المطلوبة وتدفق البيانات بين مختلف الجداول والعمليات لبناء نظام خالٍ من التضارب.")
    
    add_heading_rtl(doc, "3-2 طرق تحديد المتطلبات", level=2)
    add_body_rtl(doc, "تم تحديد المتطلبات عبر منهجيات علمية تضمنت:")
    add_body_rtl(doc, "1. المقابلات الشخصية مع مشرفي جمعيات خيرية طبية محلية في صنعاء لفهم آلية التوزيع اليدوي وعيوبها.")
    add_body_rtl(doc, "2. دراسة مجموعات الفيسبوك والواتساب المخصصة لطلب واستعارة أجهزة الأكسجين ومستلزمات العناية بالمرضى في اليمن.")
    add_body_rtl(doc, "3. حصر الجداول وقواعد البيانات اللازمة لتغطية متطلبات الرقابة والأمن الرقمي للأجهزة.")
    
    add_heading_rtl(doc, "3-3 هيكلية النظام", level=2)
    add_body_rtl(doc, "تم تقسيم النظام إلى أربعة أقسام وظيفية رئيسية (الزائر، المستفيد، المتبرع، المشرف) كما يوضح المخطط الهيكلي التالي:")
    add_image_rtl(doc, "system_hierarchy.png", "رسم توضيحي (1): مخطط هيكلية وتوزيع شاشات النظام")
    
    add_heading_rtl(doc, "3-4 فوائد التحليل", level=2)
    add_body_rtl(doc, "يضمن التحليل المنظم ما يلي:")
    add_body_rtl(doc, "1. عدم تداخل الصلاحيات (Authorization) وحماية لوحة تحكم المشرف (Admin).")
    add_body_rtl(doc, "2. منع تضارب حجز الأجهزة؛ حيث يُخفى الجهاز تلقائياً من نتائج البحث فور حجز المستفيد له وتقديمه للتقرير الطبي.")
    add_body_rtl(doc, "3. سهولة صيانة وتحديث الكود البرمجي مستقبلاً نظراً لتقسيم الوظائف بوضوح.")
    
    add_heading_rtl(doc, "3-5 نموذج الأحداث Event Modeling", level=2)
    add_body_rtl(doc, "يلخص الجدول التالي نموذج الأحداث وتفاصيل العمليات والضوابط المتبعة برمجياً:")
    
    evt_headers = ["اسم العملية", "منفذ العملية", "مدخلات العملية", "ضوابط التنفيذ والشروط", "مخرجات وتأثيرات العملية"]
    evt_rows = [
        ["تسجيل حساب جديد", "المتبرع / المستفيد", "الاسم، الهاتف، المحافظة، كلمة المرور", "عدم تكرار البريد والهاتف بقاعدة البيانات", "إنشاء حساب مع دور (Role) محدد"],
        ["إضافة جهاز طبي", "المتبرع", "اسم الجهاز، القسم، الحالة، الصور، إحداثيات الخريطة", "إدخال مدة الإعارة إجبارياً إن كانت مؤقتة", "حفظ الجهاز بحالة (معلق) لحين موافقة الإدارة"],
        ["اعتماد إعلان الجهاز", "المشرف (الأدمن)", "تحديد القرار (قبول / رفض)", "فحص ومراجعة ملاءمة الصور والأوصاف", "تحديث حالة الجهاز إلى (نشط) وظهوره في السوق"],
        ["تقديم طلب استعارة", "المستفيد", "وصف الحالة، رفع التقرير الطبي PDF/Image", "أن يكون الجهاز بحالة (نشط / متاح) حالياً", "تحديث حالة الجهاز إلى (تحت المراجعة) وإخفاؤه من البحث"],
        ["قبول طلب الاستعارة", "المشرف (الأدمن)", "تحديد القرار (قبول / رفض التقرير)", "مطابقة التقرير الطبي مع متطلبات الجهاز", "تحديث حالة الجهاز إلى (معار)، وتفعيل أزرار واتساب للطرفين"]
    ]
    add_table_rtl(doc, evt_headers, evt_rows)
    
    add_heading_rtl(doc, "3-6 مخططات تدفق البيانات DFD", level=2)
    add_heading_rtl(doc, "أولاً: مخطط السياق العام (Context Level DFD)", level=3)
    add_body_rtl(doc, "يوضح مخطط السياق العام الحدود الخارجية للنظام وكيفية تفاعل الكيانات الخارجية الثلاثة مع منصة سند:")
    add_image_rtl(doc, "context_dfd.png", "رسم توضيحي (2): مخطط السياق العام للنظام Context Level DFD")
    
    add_heading_rtl(doc, "ثانياً: مخطط تدفق البيانات التفصيلي (Level 0 DFD)", level=3)
    add_body_rtl(doc, "يفصل هذا المخطط العمليات الداخلية الأربع الرئيسية ومخازن البيانات الثلاثة المستخدمة:")
    add_image_rtl(doc, "level0_dfd.png", "رسم توضيحي (3): مخطط تدفق البيانات التفصيلي Level 0 DFD")
    
    add_heading_rtl(doc, "3-7 مخطط العلاقات الكينونية ERD", level=2)
    add_body_rtl(doc, "يوضح مخطط العلاقات الكينونية هيكل الجداول في قاعدة البيانات والعلاقات العلائقية ومفاتيح الربط والخصائص:")
    add_image_rtl(doc, "erd_diagram.png", "رسم توضيحي (4): مخطط العلاقات الكينونية لقاعدة البيانات ERD")
    
    add_heading_rtl(doc, "3-8 لغة النمذجة الموحدة UML", level=2)
    add_heading_rtl(doc, "1- مخطط حالات الاستخدام (Use Case Diagram)", level=3)
    add_body_rtl(doc, "يوضح مخطط حالات الاستخدام السيناريوهات التفاعلية للمستفيدين والمتبرعين والمسؤولين مع النظام:")
    add_image_rtl(doc, "use_case_diagram.png", "رسم توضيحي (5): مخطط حالات الاستخدام Use Case Diagram")
    
    add_heading_rtl(doc, "2- مخططات التتابع (Sequence Diagrams)", level=3)
    add_body_rtl(doc, "أ- مخطط تتابع عملية تسجيل حساب جديد:")
    add_image_rtl(doc, "sequence_register.png", "رسم توضيحي (6): مخطط تتابع عملية تسجيل الحساب")
    
    add_body_rtl(doc, "ب- مخطط تتابع عملية إدراج جهاز طبي من قِبل المتبرع:")
    add_image_rtl(doc, "sequence_add_device.png", "رسم توضيحي (7): مخطط تتابع عملية إضافة الجهاز")
    
    add_body_rtl(doc, "ج- مخطط تتابع عملية طلب استعارة جهاز ومراجعة المشرف وتفعيل الواتساب:")
    add_image_rtl(doc, "sequence_request_device.png", "رسم توضيحي (8): مخطط تتابع عملية طلب جهاز والقبول والتواصل")
    
    doc.add_page_break()

def build_chapter_4(doc):
    add_heading_rtl(doc, "الفصل الرابع: التصميم", level=1)
    
    add_heading_rtl(doc, "4-1 المقدمة", level=2)
    add_body_rtl(doc, "تركز مرحلة التصميم على تحويل النماذج التحليلية والمخططات المنطقية إلى تصاميم انسيابية وجداول تفصيلية وواجهات مستخدم جاهزة للتطبيق البرمجي.")
    
    add_heading_rtl(doc, "4-2 المخطط الانسيابي Flowcharts", level=2)
    add_body_rtl(doc, "أ- المخطط الانسيابي لإضافة جهاز طبي جديد (المتبرع):")
    add_image_rtl(doc, "flowchart_donor.png", "رسم توضيحي (9): المخطط الانسيابي لعملية إضافة جهاز")
    
    add_body_rtl(doc, "ب- المخطط الانسيابي لطلب جهاز ورفع التقرير (المستفيد):")
    add_image_rtl(doc, "flowchart_beneficiary.png", "رسم توضيحي (10): المخطط الانسيابي لعملية طلب جهاز")
    
    add_body_rtl(doc, "ج- المخطط الانسيابي لمعالجة وتدقيق الطلبات (المشرف):")
    add_image_rtl(doc, "flowchart_admin.png", "رسم توضيحي (11): المخطط الانسيابي لعملية تدقيق المشرف")
    
    add_heading_rtl(doc, "4-3 جداول قاعدة البيانات", level=2)
    add_body_rtl(doc, "تم بناء أربعة جداول علائقية رئيسية في قاعدة البيانات MySQL مصممة بمحرك InnoDB لضمان العلاقات ومفاتيح الربط التامة:")
    
    # 1. users
    add_heading_rtl(doc, "أولاً: جدول المستخدمين (users)", level=3)
    u_headers = ["اسم الحقل", "نوع البيانات", "الحجم", "المفتاح", "الوصف والخصائص"]
    u_rows = [
        ["id", "INT UNSIGNED", "10", "PRIMARY KEY", "المعرف الفريد التلقائي للمستخدم (Auto Increment)"],
        ["full_name", "VARCHAR", "100", "-", "الاسم الكامل للمستخدم (متبرع أو مستفيد)"],
        ["phone", "VARCHAR", "20", "UNIQUE", "رقم الهاتف الإلزامي للتواصل والتحقق"],
        ["email", "VARCHAR", "150", "UNIQUE", "البريد الإلكتروني المخصص لتسجيل الدخول"],
        ["password_hash", "VARCHAR", "255", "-", "كلمة المرور المشفرة بخوارزمية BCRYPT"],
        ["role", "ENUM", "-", "-", "صلاحية المستخدم ('beneficiary', 'donor', 'admin')"],
        ["governorate", "VARCHAR", "50", "-", "المحافظة الجغرافية للمستخدم (صنعاء، عدن...)"],
        ["district", "VARCHAR", "100", "-", "المديرية السكنية التابع لها المستخدم"],
        ["is_active", "TINYINT", "1", "-", "حالة تنشيط الحساب (1 نشط، 0 محظور)"],
        ["created_at", "TIMESTAMP", "-", "-", "تاريخ ووقت إنشاء الحساب تلقائياً"]
    ]
    add_table_rtl(doc, u_headers, u_rows)
    
    # 2. devices
    add_heading_rtl(doc, "ثانياً: جدول الأجهزة الطبية (devices)", level=3)
    d_headers = ["اسم الحقل", "نوع البيانات", "الحجم", "المفتاح", "الوصف والخصائص"]
    d_rows = [
        ["id", "INT UNSIGNED", "10", "PRIMARY KEY", "المعرف الفريد التلقائي للجهاز الطبي"],
        ["donor_id", "INT UNSIGNED", "10", "FOREIGN KEY", "مرتبط بالمتبرع في جدول users.id (CASCADE)"],
        ["name", "VARCHAR", "150", "-", "اسم الجهاز الطبي بدقة (مثال: مكثف أكسجين)"],
        ["category", "ENUM", "-", "-", "قسم الجهاز ('respiratory', 'mobility', 'beds_clinical', 'diagnostic')"],
        ["condition_rating", "ENUM", "-", "-", "الحالة التشغيلية للجهاز ('excellent', 'good', 'acceptable')"],
        ["description", "TEXT", "-", "-", "وصف تفصيلي للجهاز وملحقاته وحالته"],
        ["offer_type", "ENUM", "-", "-", "نوع التقديم للجمهور ('donation' تبرع، 'loan' إعارة)"],
        ["loan_duration", "VARCHAR", "50", "-", "المدة القصوى للإعارة (تظهر في حالة الإعارة)"],
        ["governorate", "VARCHAR", "50", "-", "المحافظة المتواجد بها الجهاز حالياً"],
        ["district", "VARCHAR", "100", "-", "المديرية المتواجد بها الجهاز بدقة"],
        ["latitude", "DECIMAL", "10,8", "-", "خط العرض لتحديد الموقع الجغرافي على الخريطة"],
        ["longitude", "DECIMAL", "11,8", "-", "خط الطول لتحديد الموقع الجغرافي على الخريطة"],
        ["status", "ENUM", "-", "-", "حالة الجهاز ('pending_review', 'active', 'under_request_review', 'loaned', 'rejected')"],
        ["rejection_reason", "TEXT", "-", "-", "سبب رفض الإعلان من قِبل المشرف"],
        ["admin_reviewed_by", "INT UNSIGNED", "10", "FOREIGN KEY", "المشرف الذي راجع الجهاز مرتبط بـ users.id (SET NULL)"],
        ["admin_reviewed_at", "DATETIME", "-", "-", "تاريخ مراجعة وإقرار الإعلان من المشرف"],
        ["created_at", "TIMESTAMP", "-", "-", "تاريخ ووقت إضافة الجهاز من المتبرع"]
    ]
    add_table_rtl(doc, d_headers, d_rows)
    
    # 3. device_photos
    add_heading_rtl(doc, "ثالثاً: جدول صور الأجهزة (device_photos)", level=3)
    dp_headers = ["اسم الحقل", "نوع البيانات", "الحجم", "المفتاح", "الوصف والخصائص"]
    dp_rows = [
        ["id", "INT UNSIGNED", "10", "PRIMARY KEY", "المعرف الفريد التلقائي للصورة المرفوعة"],
        ["device_id", "INT UNSIGNED", "10", "FOREIGN KEY", "مرتبط بالجهاز في جدول devices.id (CASCADE)"],
        ["file_path", "VARCHAR", "255", "-", "مسار حفظ وتخزين الصورة في السيرفر"],
        ["is_primary", "TINYINT", "1", "-", "هل الصورة هي الصورة الرئيسية للعرض (1 نعم، 0 لا)"],
        ["uploaded_at", "TIMESTAMP", "-", "-", "تاريخ ووقت رفع الصورة إلى السيرفر"]
    ]
    add_table_rtl(doc, dp_headers, dp_rows)
    
    # 4. requests
    add_heading_rtl(doc, "رابعاً: جدول طلبات الاستعارة (requests)", level=3)
    r_headers = ["اسم الحقل", "نوع البيانات", "الحجم", "المفتاح", "الوصف والخصائص"]
    r_rows = [
        ["id", "INT UNSIGNED", "10", "PRIMARY KEY", "المعرف الفريد التلقائي لطلب الحصول على الجهاز"],
        ["device_id", "INT UNSIGNED", "10", "FOREIGN KEY", "مرتبط بالجهاز المطلوب في devices.id (CASCADE)"],
        ["beneficiary_id", "INT UNSIGNED", "10", "FOREIGN KEY", "مرتبط بالمستفيد صاحب الطلب في users.id (CASCADE)"],
        ["case_description", "TEXT", "-", "-", "شرح مختصر يكتبه المريض عن حالته المرضية"],
        ["medical_doc_path", "VARCHAR", "255", "-", "مسار التقرير الطبي أو الروشتة المرفوعة للحماية"],
        ["status", "ENUM", "-", "-", "حالة الطلب لدى المشرف ('pending', 'approved', 'rejected')"],
        ["rejection_reason", "TEXT", "-", "-", "سبب رفض المشرف لطلب المريض (التقرير غير واضح مثلاً)"],
        ["admin_reviewed_by", "INT UNSIGNED", "10", "FOREIGN KEY", "المشرف الذي راجع الطلب مرتبط بـ users.id (SET NULL)"],
        ["admin_reviewed_at", "DATETIME", "-", "-", "تاريخ مراجعة واتخاذ القرار في الطلب"],
        ["created_at", "TIMESTAMP", "-", "-", "تاريخ ووقت تقديم الطلب من المريض"]
    ]
    add_table_rtl(doc, r_headers, r_rows)
    
    add_heading_rtl(doc, "4-4 واجهات المنصة", level=2)
    add_body_rtl(doc, "تم بناء وتنسيق واجهات الموقع لتكون خفيفة وعصرية بالاعتماد على تصميم responsive بالكامل لتشغيلها على الجوال. وتتلخص الواجهات في النقاط التالية:")
    add_body_rtl(doc, "1. الصفحة الرئيسية (index.php): تحتوي على شريط الإحصائيات السريعة للأجهزة المتاحة وعدد العائلات المستفيدة، وقسم توضيحي لخطوات الاستعارة والتبرع.")
    add_body_rtl(doc, "2. صفحة التسجيل والدخول (register.php, login.php): تتيح للمستخدم تحديد دوره كـ (متبرع) أو (مستفيد) وحفظ بياناته مع تطبيق حماية CSRF لسلامة البيانات.")
    add_body_rtl(doc, "3. سوق الأجهزة (marketplace.php): واجهة تفاعلية تحتوي على فلاتر المحافظات، وتصفية الأقسام الطبية ونوع الإعارة، وتعمل ديناميكياً باستخدام Vanilla JavaScript دون إعادة تحميل الصفحة.")
    add_body_rtl(doc, "4. تفاصيل الجهاز (device.php): تعرض ألبوم الصور للجهاز الطبي، حالته، مدة إعارته، موقعه الجغرافي على خريطة جوجل التفاعلية، وزر لتقديم الطلب.")
    add_body_rtl(doc, "5. لوحة تحكم المتبرع (dashboard-donor.php): تتيح له إدارة الأجهزة التي أضافها، معرفة حالتها (معلقة، نشطة، معارة)، ورؤية أسماء ومحافظات المستفيدين المقبولين لتسليمهم الأجهزة.")
    add_body_rtl(doc, "6. لوحة تحكم المستفيد (dashboard-beneficiary.php): يعرض من خلالها طلباته المعلقة، وعند موافقة المشرف، يظهر له زران تفاعليان: (اتصال هاتفي مباشر) و(تواصل فوري عبر الواتساب برابط WhatsApp API ورسالة مجهزة مسبقاً).")
    add_body_rtl(doc, "7. لوحة تحكم المشرف (admin/index.php): لوحة تحكم شاملة تعرض إحصائيات سريعة للطلبات المعلقة والأجهزة بانتظار المراجعة.")
    add_body_rtl(doc, "8. مراجعة الإعلانات والطلبات (admin/listings.php, admin/requests.php): شاشات تتيح للأدمن فحص صور الجهاز والموقع الجغرافي للموافقة عليه، أو فتح التقرير الطبي للمستفيد ودراسته طبياً لإقرار الطلب وتفعيل التواصل أو رفضه مع ذكر السبب.")
    
    doc.add_page_break()

def build_chapter_5(doc):
    add_heading_rtl(doc, "الفصل الخامس: الخاتمة والتوصيات", level=1)
    
    add_heading_rtl(doc, "5-1 الخاتمة", level=2)
    add_body_rtl(doc, "في ختام صفحات هذا المشروع البحثي والتطبيقي، نؤكد أن منصة (سَنَد) تمثل نموذجاً عملياً وفعالاً لحل مشكلة إنسانية واجتماعية واضحة تلمس الأسر اليمنية بشكل يومي. فبدلاً من هدر الموارد الطبية أو بقائها مخزنة بشكل خامل، أثبتت المنصة أن استخدام التكنولوجيا التكافلية يبسط التواصل المباشر ويسرع إنقاذ الأرواح.")
    add_body_rtl(doc, "كما أن البناء الفني للموقع بالاعتماد على PHP الخام وMySQL بدون أطر عمل جاهزة ساهم في بناء أسس برمجية متينة للمطورين، تعكس إمكاناتهم في تأمين معالجة الملفات والتعامل مع جلسات المستخدمين وحماية قاعدة البيانات والخصوصية الطبية.")
    
    add_heading_rtl(doc, "5-2 التوصيات", level=2)
    add_body_rtl(doc, "يوصي الباحث ببعض التوصيات الضرورية لضمان استمرارية المنصة وتطويرها:")
    add_body_rtl(doc, "1. إقامة شراكات رسمية مع الجمعيات الطبية والخيرية اليمنية لتبني المنصة وتوفير لجان إشراف طبية معتمدة لمراجعة التقارير.")
    add_body_rtl(doc, "2. توعية المجتمع ونشر ثقافة التبرع وإعارة الأجهزة الطبية المستعملة عبر وسائل التواصل لزيادة عدد المتبرعين.")
    add_body_rtl(doc, "3. تشجيع المبادرات الشبابية الطوعية لتقديم خدمات صيانة الأجهزة الطبية التبرعية للتأكد من سلامتها التشغيلية قبل الإعارة.")
    
    add_heading_rtl(doc, "5-3 العمل المستقبلي للتطوير", level=2)
    add_body_rtl(doc, "يتطلع المشروع إلى التطويرات المستقبلية التالية:")
    add_body_rtl(doc, "1. بناء وتطوير تطبيقات هاتف ذكي (Native Android & iOS Apps) لسهولة التصفح ورفع الصور والموقع.")
    add_body_rtl(doc, "2. دمج الذكاء الاصطناعي (AI) لفحص وتصنيف التقارير الطبية المرفوعة تلقائياً للتحقق من مصداقيتها وتقليل العبء على المشرفين.")
    add_body_rtl(doc, "3. دمج نظام تتبع توصيل للأجهزة بالتعاون مع مبادرات النقل التطوعي أو شركات التوصيل المحلية لتسهيل النقل للمرضى العاجزين.")
    
    add_heading_rtl(doc, "5-4 المعوقات", level=2)
    add_body_rtl(doc, "واجه فريق تطوير المشروع بعض العقبات أثناء التنفيذ والتخطيط:")
    add_body_rtl(doc, "1. ضعف سرعات الإنترنت والاتصال بالشبكة المحلية مما أثر على فحص وتكامل خرائط جوجل.")
    add_body_rtl(doc, "2. انقطاع التيار الكهربائي المتكرر مما أثر على استقرار بيئة التطوير المحتضنة (XAMPP).")
    add_body_rtl(doc, "3. صعوبة صياغة معايير طبية صارمة للتحقق البشري من صحة المستندات دون وجود طبيب مختص دائم.")
    
    add_heading_rtl(doc, "5-5 المراجع", level=2)
    p_ref = doc.add_paragraph()
    make_paragraph_rtl(p_ref)
    add_arabic_run(p_ref, "1. مؤسسة MedShare الدولية للتبرع بالأجهزة الطبية: www.medshare.org\n", size=10)
    add_arabic_run(p_ref, "2. شبكة Freecycle العالمية لإعادة الاستخدام والتبادل: www.freecycle.org\n", size=10)
    add_arabic_run(p_ref, "3. دليل هندسة البرمجيات وبناء قواعد البيانات العيانية، الطبعة الخامسة، روجر برسمن.\n", size=10)
    add_arabic_run(p_ref, "4. توثيق لغة البرمجة PHP الرسمية وطرق تأمين PDO Prepared Statements: www.php.net\n", size=10)
    add_arabic_run(p_ref, "5. إحصائيات القطاع الصحي وأثر الأزمة الاقتصادية في اليمن، منظمة الصحة العالمية WHO 2024.\n", size=10)
    
    # End page (Thank you / Calligraphy)
    doc.add_page_break()
    for i in range(10):
        doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_arabic_run(p_end, "تم بحمد الله وتوفيقه", bold=True, size=24, color=RGBColor(0, 76, 153))

def main():
    print("Creating document skeleton...")
    doc = create_document()
    
    print("Building Cover Page...")
    build_cover_page(doc)
    
    print("Building Preliminary Pages...")
    build_preliminary_pages(doc)
    
    print("Building Chapter 1...")
    build_chapter_1(doc)
    
    print("Building Chapter 2...")
    build_chapter_2(doc)
    
    print("Building Chapter 3...")
    build_chapter_3(doc)
    
    print("Building Chapter 4...")
    build_chapter_4(doc)
    
    print("Building Chapter 5...")
    build_chapter_5(doc)
    
    print("Adding header and footer...")
    add_header_footer(doc)
    
    # Save the document
    output_docx = os.path.join(os.path.dirname(__file__), "sanad_documentation.docx")
    doc.save(output_docx)
    print(f"Document saved successfully as: {output_docx}")

if __name__ == "__main__":
    main()
